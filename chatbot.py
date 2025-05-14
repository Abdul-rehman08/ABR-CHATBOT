# import streamlit as st
# import google.generativeai as genai

# # --- CONFIGURE API ---
# genai.configure(api_key="AIzaSyA5fU4CV5mfbk5pYleH176m_tHmpLp1XQY")
# model = genai.GenerativeModel('gemini-1.5-flash')

# # --- PAGE SETTINGS ---
# st.set_page_config(page_title="ABRGPT", page_icon="🤖", layout="wide")
# st.markdown(
#     "<style>html, body, [class*='css']  { font-family: 'Segoe UI', sans-serif; }</style>",
#     unsafe_allow_html=True
# )

# # --- CUSTOM STYLING ---
# st.markdown("""
#     <style>
#         .main {
#             # background: linear-gradient(135deg, #e3f2fd 0%, #ffffff 100%);
#         }
#         .block-container {
#             padding: 2rem 3rem;
#         }
#         .chat-bubble-user {
#             background-color: #DCF8C6;
#             color: black;
#             padding: 0.75rem 1rem;
#             border-radius: 1.5rem;
#             max-width: 75%;
#             margin-left: auto;
#             margin-bottom: 1rem;
#         }
#         .chat-bubble-bot {
#             background-color: #F1F0F0;
#             color: black;
#             padding: 0.75rem 1rem;
#             border-radius: 1.5rem;
#             max-width: 75%;
#             margin-right: auto;
#             margin-bottom: 1rem;
#         }
#         footer {visibility: hidden;}
#     </style>
# """, unsafe_allow_html=True)

# # --- SIDEBAR ---
# with st.sidebar:
#     # st.image("https://upload.wikimedia.org/wikipedia/commons/5/5e/Google_Gemini_logo.svg", width=200)
#     st.image("cabr.png", width=350)
#     st.markdown("### 🤖 ABRGPT")
#     st.markdown("Ask anything & get instant responses from 'ABRGPT AI'.")
#     if st.button("🧹 New Chat"):
#         st.session_state.chat = model.start_chat(history=[])
#         st.rerun()

#     if "chat" in st.session_state:
#         history_text = "\n".join(
#             f"{msg.role}: {msg.parts[0].text}" for msg in st.session_state.chat.history
#         )
#         st.download_button("💾 Download Chat", history_text, file_name="gemini_chat.txt")

# # --- INIT CHAT ---
# if "chat" not in st.session_state:
#     st.session_state.chat = model.start_chat(history=[])

# st.markdown("<h2 style='text-align: center;'>🤖 ABRGPT Chat Assistant</h2>", unsafe_allow_html=True)

# # --- CHAT DISPLAY ---
# for msg in st.session_state.chat.history:
#     bubble_class = "chat-bubble-user" if msg.role == "user" else "chat-bubble-bot"
#     st.markdown(f"<div class='{bubble_class}'>{msg.parts[0].text}</div>", unsafe_allow_html=True)

# # --- USER INPUT ---
# prompt = st.chat_input("Type your message...")

# if prompt:
#     st.markdown(f"<div class='chat-bubble-user'>{prompt}</div>", unsafe_allow_html=True)

#     with st.spinner("Gemini is thinking..."):
#         response = st.session_state.chat.send_message(prompt)
#         st.markdown(f"<div class='chat-bubble-bot'>{response.text}</div>", unsafe_allow_html=True)

import streamlit as st
import google.generativeai as genai
import pdfplumber
import docx
from docx import Document
from fpdf import FPDF
import io

# --- CONFIGURE API ---
genai.configure(api_key="YOUR_API_KEY")
model = genai.GenerativeModel('gemini-1.5-flash')

# --- PAGE SETTINGS ---
st.set_page_config(page_title="ABRGPT", page_icon="🤖", layout="wide")

# --- INIT SESSION STATE ---
if "chat" not in st.session_state:
    st.session_state.chat = model.start_chat(history=[])
if "file_context" not in st.session_state:
    st.session_state.file_context = ""
if "files_uploaded" not in st.session_state:
    st.session_state.files_uploaded = []

# --- SIDEBAR ---
with st.sidebar:
    st.image("cabr.png", width=350)
    st.markdown("### 🤖 ABRGPT")
    if st.button("🧹 New Chat"):
        st.session_state.chat = model.start_chat(history=[])
        st.session_state.file_context = ""
        st.session_state.files_uploaded = []
        st.rerun()

    if st.session_state.chat.history:
        history_text = "\n".join(
            f"{msg.role}: {msg.parts[0].text}" for msg in st.session_state.chat.history
        )
        st.download_button("💾 Download Chat", history_text, file_name="gemini_chat.txt")

# --- FILE UPLOAD ---
uploaded_files = st.file_uploader("📎 Upload documents (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"], accept_multiple_files=True)

def extract_text(file):
    text = ""
    if file.type == "application/pdf":
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ""
    elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = docx.Document(file)
        text = "\n".join([para.text for para in doc.paragraphs])
    elif file.type == "text/plain":
        text = file.read().decode("utf-8")
    return text

if uploaded_files:
    combined_text = ""
    for file in uploaded_files:
        file_text = extract_text(file)
        st.session_state.files_uploaded.append({
            "name": file.name,
            "type": file.type,
            "text": file_text
        })
        combined_text += f"\n\n--- {file.name} ---\n\n{file_text}"

    st.session_state.file_context = combined_text
    st.success("✅ Files uploaded and processed!")

    with st.expander("📄 View Combined File Content"):
        st.text_area("File Content", combined_text, height=250)

# --- CONVERSION FUNCTION ---
def convert_text_to_file(text, format):
    buf = io.BytesIO()
    filename = f"converted_file.{format}"

    if format == "txt":
        buf.write(text.encode("utf-8"))
    elif format == "docx":
        doc = Document()
        for line in text.split("\n"):
            doc.add_paragraph(line)
        doc.save(buf)
    elif format == "pdf":
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Arial", size=12)
        for line in text.split("\n"):
            pdf.cell(200, 10, txt=line.encode('latin-1', errors='ignore').decode('latin-1'), ln=True)
        pdf.output(buf)
    buf.seek(0)
    return filename, buf

# --- FUNCTION TO TRIM LONG TEXT ---
MAX_CHARS = 20000  # Keep safely under Gemini input token limit

def trim_text(text, max_chars):
    return text if len(text) <= max_chars else text[-max_chars:]

# --- CHAT INTERFACE ---
st.markdown("<h2 style='text-align: center;'>🤖 ABRGPT Chat Assistant</h2>", unsafe_allow_html=True)

for msg in st.session_state.chat.history:
    bubble_class = "chat-bubble-user" if msg.role == "user" else "chat-bubble-bot"
    st.markdown(f"<div class='{bubble_class}'>{msg.parts[0].text}</div>", unsafe_allow_html=True)

prompt = st.chat_input("Type your message...")

if prompt:
    st.markdown(f"<div class='chat-bubble-user'>{prompt}</div>", unsafe_allow_html=True)

    # Check for conversion commands
    if "convert" in prompt.lower():
        if "pdf" in prompt.lower():
            file_format = "pdf"
        elif "docx" in prompt.lower() or "word" in prompt.lower():
            file_format = "docx"
        elif "txt" in prompt.lower() or "text" in prompt.lower():
            file_format = "txt"
        else:
            file_format = None

        if file_format:
            filename, converted_file = convert_text_to_file(st.session_state.file_context, file_format)
            st.download_button(f"⬇️ Download as {file_format.upper()}", data=converted_file, file_name=filename)
            st.markdown(f"<div class='chat-bubble-bot'>✅ File converted to {file_format.upper()}!</div>", unsafe_allow_html=True)
        else:
            st.markdown(f"<div class='chat-bubble-bot'>❌ Could not detect file format. Try saying 'Convert to PDF', 'Make DOCX', etc.</div>", unsafe_allow_html=True)
    else:
        # Trim the file content if it's too long
        if st.session_state.file_context:
            trimmed_context = trim_text(st.session_state.file_context, MAX_CHARS)
            full_prompt = f"{trimmed_context}\n\nUser Question: {prompt}"
        else:
            full_prompt = prompt

        with st.spinner("Gemini is thinking..."):
            response = st.session_state.chat.send_message(full_prompt)
            st.markdown(f"<div class='chat-bubble-bot'>{response.text}</div>", unsafe_allow_html=True)

